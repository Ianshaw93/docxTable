from pathlib import Path
import datetime
import os

from docx import Document
from docx.shared import Pt

# NB: document string should have no spaces or be a r'' string
document_name = "Template-Colleges.docx"
document_path = Path(__file__).parent /document_name


def Delete_row_in_table(table, row):
    if type(table) == int:
        table_object = document.tables[table]
    else:
        table_object = table
    table_object._tbl.remove(table_object.rows[row]._tr)

today = datetime.datetime.today()

output_path = Path(__file__).parent / "outputReports"/"report.docx"

document = Document(document_name)



def replace_table_cell_content(cell, replacement_text, is_bold=False, alignment=1):
    # later make object with fonts etc
    cell.text = replacement_text
    paragraphs = cell.paragraphs
    paragraphs[0].alignment = alignment # 1 = centered
    # for run in paragraphs.runs:
    run = paragraphs[0].runs
    font = run[0].font
    font.size = Pt(9) # pull from object
    font.name = "Arial"
    # font.color.rgb = RGBColor(64,64,64) #gray color # to be pulled from object
    font.bold = is_bold

def alter_table_rows(total_rows, table, header_rows = 1):
    current_rows = len(table.rows)
    required_rows = total_rows + header_rows
    # remove bottom rows
    rows_to_remove = current_rows - required_rows
    for i in range(rows_to_remove):
        Delete_row_in_table(table, row=-1)

def locate_table_from_cell(cell_row_index, cell_column_index, cell_text):
    table_list = []
    
    for table in document.tables:
        cell = table.rows[cell_row_index].cells[cell_column_index]
        if cell_text in cell.text:
            table_list.append(table)
    return table_list

college_table = locate_table_from_cell(cell_row_index=0, cell_column_index=0, cell_text="College")[0]
repo_file = "colleges_data.csv"

with open(repo_file, "r+", encoding="utf8") as f:
    repo_list = f.readlines()[1:]
repo_list = [f.strip("\n") for f in repo_list]
split_repo_list = [f.split(",") for f in repo_list]


# additional row for Total
table_rows = len(split_repo_list)+1
alter_table_rows(total_rows=table_rows, table=college_table)

changes_total = 0
for row_index in range(1, table_rows+1): # skip heading_row and incorporate all additional rows - range function does not include max
    columns = len(college_table.rows[0].cells)
    for column_index in range(columns):
        current_cell = college_table.rows[row_index].cells[column_index]

        if row_index == (table_rows):
            # total row
            if column_index > 0: # needs total not from data for other columns
                # add totals from column each time
                # loop through rows -> add column to total
                if column_index == 3:
                    current_text = str(changes_total)
                else:
                    column_total = 0
                    for row_i in range(1, table_rows):
                        # if row_i > 0 and row_i is not table_rows:
                        current_row_text = split_repo_list[row_i-1]
                        current_row_col_text = current_row_text[column_index]
                        column_total += int(current_row_col_text)
                    current_text = str(column_total)
                replace_table_cell_content(cell=current_cell, replacement_text=current_text, is_bold=False, alignment=1)
            else: 
                replace_table_cell_content(cell=current_cell, replacement_text="Total", is_bold=True, alignment=1)

                
        elif row_index > 0: # should not alter row zero
            current_text_row = split_repo_list[row_index-1]
            if column_index <= 2:
                current_text = current_text_row[column_index]
            if column_index == 3:
                current_change = int(current_text_row[column_index-2]) - int(current_text_row[column_index-1])
                changes_total += current_change
                current_text = str(current_change)
            if column_index == 0:
                is_bold = True
            else:
                is_bold = False
            replace_table_cell_content(cell=current_cell, replacement_text=current_text, is_bold=is_bold, alignment=1)

document.save(output_path)
os.startfile(output_path)
# loop through rows
# loop through cells
# make column 0 bold
# make final row bold
